"""生成「专利检索智能体汇报PPT讲稿」Word 文档 (python-docx 版本).

特点:
- 与 14 页 PPT 一一对应, 每段含页码标注、动作提示、口播讲稿
- 含封面、使用说明、时间分配总览表、附录(节奏建议)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============== 配色 ==============
PRIMARY = RGBColor(0x38, 0x59, 0xFF)
DARK = RGBColor(0x1C, 0x20, 0x38)
TEXT = RGBColor(0x2A, 0x2D, 0x3A)
LIGHT = RGBColor(0x6B, 0x70, 0x80)
ACCENT = RGBColor(0xFF, 0x6B, 0x35)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_CN = "Microsoft YaHei"

OUTPUT = r"D:\Claude Code Files\Project_Patent search system_v1\专利检索智能体汇报PPT讲稿.docx"


# ============== 工具函数 ==============
def set_cell_bg(cell, hex_color):
    """设置单元格背景色."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_run(p, text, size=11, bold=False, color=TEXT, italic=False, font=FONT_CN):
    """添加 run 并设置中英文兼容字体."""
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    return run


def add_para(doc, runs=None, text=None, size=11, bold=False, color=TEXT,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.5,
             italic=False):
    """添加段落 (支持单文本或多 run)."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if runs is None:
        runs = [{"text": text, "size": size, "bold": bold, "color": color, "italic": italic}]
    for r in runs:
        add_run(p, r["text"], size=r.get("size", size), bold=r.get("bold", False),
                color=r.get("color", TEXT), italic=r.get("italic", False))
    return p


def add_page_tag(doc, page_no, page_title, duration):
    """PPT 页码标注条."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, f"【PPT 第 {page_no} 页】", size=12, bold=True, color=PRIMARY)
    add_run(p, f"  {page_title}", size=12, bold=True, color=DARK)
    add_run(p, f"   (时长约 {duration})", size=10, color=LIGHT, italic=True)
    # 底部彩色边框
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "3859FF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_action_hint(doc, text):
    """动作提示行."""
    add_para(doc, runs=[
        {"text": "\U0001F3AC 动作提示: ", "size": 10, "bold": True, "color": ACCENT},
        {"text": text, "size": 10, "color": LIGHT, "italic": True},
    ], before=2, after=8)


def add_speech(doc, text):
    """口播讲稿行."""
    add_para(doc, runs=[
        {"text": "\U0001F5E3  ", "size": 11},
        {"text": text, "size": 11, "color": TEXT},
    ], before=2, after=6, line=1.6)


def add_note(doc, text):
    """补充说明行."""
    add_para(doc, runs=[
        {"text": "\U0001F4A1 ", "size": 10},
        {"text": text, "size": 10, "color": LIGHT, "italic": True},
    ], before=2, after=8, line=1.5)


def add_emphasis(doc, label, content):
    """强调行 (粗体标签 + 内容)."""
    add_para(doc, runs=[
        {"text": label, "size": 11, "bold": True, "color": ACCENT},
        {"text": " " + content, "size": 11, "color": TEXT},
    ], before=2, after=6)


def add_heading(doc, text, level=1):
    """添加标题 (用样式)."""
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(12)
        add_run(p, text, size=18, bold=True, color=DARK)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        add_run(p, text, size=14, bold=True, color=PRIMARY)
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        add_run(p, text, size=12, bold=True, color=DARK)
    return p


# ============== 文档构建 ==============
def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = FONT_CN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)

    # 页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ============== 封面 ==============
    add_para(doc, runs=[{"text": "专利检索智能体", "size": 28, "bold": True, "color": DARK}],
             align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=6)
    add_para(doc, runs=[{"text": "AI 辅助专利审查系统 \u2014 项目阶段汇报讲稿",
                          "size": 14, "color": PRIMARY}],
             align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

    # 信息表
    info_rows = [
        ("汇报人", "周夏恒 (医化部 \u00B7 大模型应用工作组)"),
        ("汇报对象", "部门领导"),
        ("汇报时长", "约 10 分钟 (含视频演示 2 分 38 秒)"),
        ("配套文件", "专利检索智能体汇报PPT.pptx (14 页)"),
        ("日期", "2026 年 6 月"),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (k, v) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(9)
        # 左列
        c0 = row.cells[0]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(c0, "F5F7FA")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p0, k, size=11, bold=True, color=DARK)
        # 右列
        c1 = row.cells[1]
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        add_run(p1, v, size=11, color=TEXT)

    # ============== 使用说明 ==============
    add_heading(doc, "使用说明", level=2)
    add_para(doc, text="本讲稿对应 14 页 PPT, 按汇报顺序逐页给出: \u2460 动作提示 (在 PPT 上做什么操作); \u2461 口播讲稿 (直接照读或转述)。每段口播稿都标注了对应的 PPT 页码与建议时长, 便于您在排练时把控节奏。",
             after=8)
    add_para(doc, text="总时长控制在 10 分钟左右。如果时间紧张, 可优先压缩第 13 页 (后续规划) 与第 14 页 (收尾)。视频段落 (第 5 页) 建议提前预演, 避免现场操作卡顿。",
             after=8)
    add_para(doc, runs=[
        {"text": "排练建议: ", "size": 11, "bold": True, "color": ACCENT},
        {"text": "全文大声朗读一遍约需 9 分 30 秒, 留 30 秒缓冲用于切换页面与视频播放。", "size": 11},
    ], after=12)

    # ============== 时间分配总览表 ==============
    add_heading(doc, "时间分配总览", level=2)
    time_rows = [
        ["时段", "PPT 页", "章节", "时长"],
        ["0:00 \u2013 0:40", "第 1\u20132 页", "开场 (封面 + 议程)", "40 秒"],
        ["0:40 \u2013 1:30", "第 3\u20134 页", "项目背景与目标", "50 秒"],
        ["1:30 \u2013 4:10", "第 5 页", "系统宣传片 (视频演示)", "2 分 40 秒"],
        ["4:10 \u2013 6:40", "第 6\u20139 页", "方案介绍 (架构/能力/流程/机制)", "2 分 30 秒"],
        ["6:40 \u2013 9:00", "第 10\u201312 页", "成果演示 (看板/报告/质量)", "2 分 20 秒"],
        ["9:00 \u2013 9:40", "第 13 页", "项目现状与后续规划", "40 秒"],
        ["9:40 \u2013 10:00", "第 14 页", "总结收尾 + Q&A 引子", "20 秒"],
    ]
    table = doc.add_table(rows=len(time_rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.5), Cm(3.5), Cm(6.5), Cm(3)]
    for i, row_data in enumerate(time_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
            if i == 0:
                set_cell_bg(cell, "3859FF")
                add_run(p, cell_text, size=11, bold=True, color=WHITE)
            else:
                add_run(p, cell_text, size=11, color=TEXT)

    doc.add_page_break()

    # ============== 逐页讲稿 ==============
    add_heading(doc, "逐页讲稿", level=1)

    # --- 第 1 页 封面 ---
    add_page_tag(doc, 1, "封面", "15 秒")
    add_action_hint(doc, "PPT 打开后停留 3 秒, 让领导看清标题, 再开口。")
    add_speech(doc, "各位领导好, 今天向各位汇报我自主研发的\u201C专利检索智能体\u201D项目。"
                    "这是一个面向专利审查员实际工作场景的 AI 辅助审查系统, 聚焦\u201C高效、精准、可追溯\u201D三个核心目标。"
                    "下面我从方案设计和成果演示两个方面进行汇报。")
    add_note(doc, "封面信息: 标题\u201C专利检索智能体\u201D, 副标题\u201CAI 辅助专利审查系统\u2014 方案设计与成果演示\u201D, 汇报人周夏恒, 部门医化部 \u00B7 大模型应用工作组。")

    # --- 第 2 页 议程 ---
    add_page_tag(doc, 2, "汇报议程", "15 秒")
    add_action_hint(doc, "逐项点出 4 个章节, 每项停顿一下。")
    add_speech(doc, "本次汇报分为四个部分: 第一部分介绍项目的研发背景与目标; "
                    "第二部分从系统架构、工作流程、核心能力三个维度展开方案介绍; "
                    "第三部分通过实际看板、报告样例和操作录屏展示项目成果; "
                    "最后简要说明项目现状与后续规划。")

    # --- 第 3 页 背景 ---
    add_page_tag(doc, 3, "项目背景: 传统检索的痛点", "25 秒")
    add_action_hint(doc, "四个痛点卡片依次点出, 配合手势。")
    add_speech(doc, "在日常审查工作中, 我深刻体会到现有检索工具存在四类突出问题。"
                    "第一, 经验依赖\u2014\u2014现有工具高度依赖审查员个人经验, 新审查员上手周期长、检索质量参差不齐。"
                    "第二, 覆盖有限\u2014\u2014单一引擎覆盖度有限, 跨语种、跨领域的相关文献容易遗漏。"
                    "第三, 汇总繁重\u2014\u2014多份对比文献需要人工逐条阅读、汇总去重, 工作量极大。"
                    "第四, 追溯困难\u2014\u2014检索结果缺乏结构化呈现, 引用追溯不便。"
                    "这些痛点正是本项目要解决的核心问题。")
    add_emphasis(doc, "底部金句:", "\u201C让审查员从机械检索中解放出来, 专注于创造性判断\u201D\u2014\u2014这是项目的核心使命。")

    # --- 第 4 页 目标 ---
    add_page_tag(doc, 4, "项目目标", "15 秒")
    add_action_hint(doc, "三大目标卡片从左到右依次点出。")
    add_speech(doc, "围绕这些痛点, 项目设定三大目标。"
                    "第一, 提效\u2014\u2014大幅压缩单件专利的对比文件检索耗时。"
                    "第二, 提质\u2014\u2014通过多模型并行检索, 让覆盖度优于任何单一引擎。"
                    "第三, 可追溯\u2014\u2014每条对比文献都具备完整的来源链路, 便于审查员直接引用。"
                    "这三大目标, 既是业务诉求, 也是技术挑战。")

    # --- 第 5 页 视频 ---
    add_page_tag(doc, 5, "系统宣传片: 核心价值主张", "2 分 40 秒 (含视频 2 分 38 秒)")
    add_action_hint(doc, "切到第 5 页后, 点击视频开始播放; 播放期间不要讲解, 与领导一起观看; 视频结束后接下面这句过渡。")
    add_speech(doc, "下面请观看系统的完整操作录屏, 时长约 2 分 38 秒。"
                    "录屏中完整呈现了从登录、上传专利、AI 自动解析、多模型配置、并行检索执行, "
                    "到实时看板监控、报告自动生成与导出的全流程。")
    add_emphasis(doc, "【视频结束后过渡】", "视频中展示的是系统当前已落地的完整能力。接下来, 我用几页 PPT 拆解背后的方案设计。")
    add_note(doc, "视频文件位于项目根目录: 专利检索智能体操作录屏.mp4 (485 MB)。如现场无法播放, 可改为口述介绍要点, 或请领导会后再看。")

    # --- 第 6 页 架构 ---
    add_page_tag(doc, 6, "系统架构: 三层 + 多模型适配层", "50 秒")
    add_action_hint(doc, "从下往上指三层架构, 再指右侧模型适配层。")
    add_speech(doc, "系统采用经典的\u201C三层 + 适配层\u201D架构。"
                    "最下层是数据层, 基于 Supabase 平台, 提供 Postgres 数据库、文件存储、身份认证和实时推送四项能力。"
                    "中间是业务层, 是一个独立的 Node.js Worker 进程, 通过 pg-boss 队列负责任务调度、AI 调度、报告生成和失败重试。"
                    "最上层是表现层, 基于 Next.js 16 加 React 19 实现, 提供用户交互、流程看板、报告查看等界面。")
    add_speech(doc, "在三层之外, 我们设计了一套多 AI 引擎适配层。"
                    "目前已经接入了 6 个数据源, 分别是 Kimi、智谱、通义千问、DeepSeek、秘塔, 以及通用的 OpenAI 兼容接口。"
                    "每个引擎都有独立的适配器实现, 新增模型的接入成本几乎为零。")

    # --- 第 7 页 核心能力 ---
    add_page_tag(doc, 7, "六大核心能力 (含 6 大模型 LOGO)", "50 秒")
    add_action_hint(doc, "先指顶部 6 个 LOGO 横排, 再依次介绍下方 6 个能力卡片 (每张卡片一句话带过)。")
    add_speech(doc, "顶部展示了系统已集成的 6 大 AI 引擎, 从左到右依次是 Kimi、智谱、通义千问、DeepSeek、秘塔和 OpenAI 兼容接口, "
                    "全部采用可插拔适配方式接入。")
    add_speech(doc, "在功能层面, 系统具备六大核心能力。"
                    "第一, 多模型并行检索\u2014\u2014通过模型与策略的笛卡尔积展开, 让一次检索覆盖多个维度的相关文献。"
                    "第二, 可插拔 AI 适配器\u2014\u2014统一接口屏蔽各厂商差异。"
                    "第三, 可视化流程看板\u2014\u2014基于 React Flow 实时呈现解析、检索、报告生成的全链路。"
                    "第四, 智能去重与排序\u2014\u2014先按 URL 归一化、再按标题归一化去重, 并由独立的报告模型完成 Top-N 选优。"
                    "第五, 质量感知与告警\u2014\u2014对缺失作者、日期、链接的文献自动标注。"
                    "第六, 可追溯审计\u2014\u2014每篇文献携带来源平台、检索策略和任务 ID, 支持 GB/T 7714 引用格式输出。")

    # --- 第 8 页 工作流 ---
    add_page_tag(doc, 8, "审查员工作流: 六步完成检索", "50 秒")
    add_action_hint(doc, "横向依次点出 6 个步骤方框; 底部可指登录截图。")
    add_speech(doc, "审查员的实际使用流程非常简洁, 六步即可完成检索。"
                    "第一步, 上传专利\u2014\u2014支持 PDF、DOCX、XLSX、TXT 任一格式。"
                    "第二步, AI 解析\u2014\u2014系统自动抽取技术主题、申请人、权利要求、核心发明点等结构化字段。"
                    "第三步, 配置检索\u2014\u2014配置模型与策略的检索矩阵, 系统按经验预填默认值, 审查员可微调。"
                    "第四步, 启动并行\u2014\u2014后台自动执行 N\u00D7M 个子任务。"
                    "第五步, 流程看板\u2014\u2014React Flow 实时显示每个子任务的状态。"
                    "第六步, 生成报告\u2014\u2014Top-N 报告自动生成, 支持 Markdown 和 DOCX 导出, 并允许审查员人工评级与备注。")
    add_emphasis(doc, "下方三项关键能力:", "灵活的文件处理 (全格式支持)、智能 Prompt 编辑 (按经验定制)、人机协作闭环 (报告可改可评)。")

    # --- 第 9 页 笛卡尔积 ---
    add_page_tag(doc, 9, "核心机制: 模型 \u00D7 策略 笛卡尔积并行", "50 秒")
    add_action_hint(doc, "指左下角 3\u00D73 矩阵, 强调\u201C9 个子任务同时执行\u201D; 再指右侧关键设计。")
    add_speech(doc, "这一页是整个方案最核心的机制\u2014\u2014模型与策略的笛卡尔积并行。"
                    "举一个实际例子: 如果审查员选择了 3 个模型和 3 个策略, 系统会自动展开为 9 个检索子任务同时执行。")
    add_speech(doc, "关键设计有四点: "
                    "第一, 单模型并发上限设为 2, 避免触发厂商限流; "
                    "第二, 失败任务自动重试, 采用指数退避策略; "
                    "第三, 全局设置 20 分钟超时保护; "
                    "第四, 全程支持审查员取消操作。"
                    "所有子任务的结果汇入统一的去重与排序模块, 由独立的报告模型完成 Top-N 选优。")
    add_emphasis(doc, "底部价值金句:", "这种\u201C集思广益\u201D式检索的覆盖度, 是任何单引擎检索工具无法企及的\u2014\u2014多源印证, 既广又准。")

    # --- 第 10 页 看板 ---
    add_page_tag(doc, 10, "成果演示一: 实时流程看板", "1 分钟")
    add_action_hint(doc, "放大看板截图, 重点指 3 列模型 \u00D7 3 行策略的节点布局, 以及右侧\u201C对比文献 38 篇\u201D。")
    add_speech(doc, "下面进入成果演示。第一项是实时流程看板。"
                    "这是基于 React Flow 实现的可视化界面, 每个子任务对应一个节点, 状态实时反映后端执行进度。"
                    "可以看到四种状态: 待执行、执行中、成功、失败。")
    add_speech(doc, "以这次真实运行为例, 待审专利是 CN118278483A, 使用了智谱 GLM-5.1、秘塔 AI、Kimi K2.6 三个模型, "
                    "每个模型展开追踪检索、主要技术方案步骤检索、发明构思检索三种策略。"
                    "整个任务从 22:03 启动, 22:09 完成, 总耗时 6 分 30 秒, 最终汇聚出 38 篇对比文献。")
    add_emphasis(doc, "亮点强调:", "无需前端轮询, 基于 Supabase Realtime 推送; 支持随时取消; 排队位置可见。")

    # --- 第 11 页 报告样例 ---
    add_page_tag(doc, 11, "成果演示二: 报告样例 (前 3 篇)", "1 分钟")
    add_action_hint(doc, "左侧报告样例截图, 右侧关键特性卡片, 逐项对应。")
    add_speech(doc, "第二项成果是结构化报告。"
                    "这是一份真实运行生成的报告样例, 待审专利是 CN118429689A, 主题是基于注意力频域 GAN 的对抗样本生成方法。"
                    "系统自动检索并汇总了 10 篇相关对比文献, 涵盖 Google Patents 的同族专利和 arXiv 的学术论文, 按相关度排序输出。")
    add_speech(doc, "每篇文献都包含完整的结构化字段: 标题、来源平台、作者、发表时间、文献链接、相关描述。"
                    "以第一篇为例, 来源是 DeepSeek \u00D7 追踪检索, 直接命中了发明人完全一致的同族专利 CN114510724A, 相关度极高。")
    add_emphasis(doc, "右侧关键特性:", "结构化字段、来源链路 (来源平台 \u00D7 检索策略)、可追溯 (任务 ID)、标准化引用 (GB/T 7714)、多格式导出 (Markdown + DOCX)。")

    # --- 第 12 页 质量感知 ---
    add_page_tag(doc, 12, "成果演示三: 质量感知与可追溯性", "1 分钟")
    add_action_hint(doc, "左侧 4 条质量规则, 右侧 5 篇样例文献; 指第 6 篇\u201C标题缺失\u201D作为典型。")
    add_speech(doc, "第三项成果是质量感知与可追溯性, 这是系统区别于普通检索工具的关键。"
                    "系统对每篇对比文献都进行质量评估, 自动暴露数据缺陷。")
    add_speech(doc, "四条质量规则: "
                    "第一, 缺失字段告警\u2014\u2014若日期、链接、描述缺失, 系统如实标注, 不做虚假填充; "
                    "第二, 作者未知标注\u2014\u2014若作者字段为\u201C未知\u201D, 保留原始状态而非编造; "
                    "第三, 质量分过滤\u2014\u2014每篇文献计算 0-100 的质量分, 低于阈值 (默认 50) 的直接过滤; "
                    "第四, 疑似重复识别\u2014\u2014同族专利识别为相关但保留为独立条目, 供审查员判断。")
    add_speech(doc, "右侧展示了 5 篇真实文献的来源链路样例。"
                    "可以看到, 第 6 篇标题为\u201C文献1\u201D是数据缺陷的典型, 系统如实保留了它的不完整状态。"
                    "这种\u201C诚实暴露\u201D的方式, 让审查员对数据可信度有清晰判断。")

    # --- 第 13 页 规划 ---
    add_page_tag(doc, 13, "项目现状与后续规划", "40 秒")
    add_action_hint(doc, "左侧已完成能力列表快速过; 重点在右侧两个方向, 请领导指示。")
    add_speech(doc, "项目目前已完成检索流程的全链路闭环, "
                    "包括上传、解析、检索、报告生成、导出等核心环节; "
                    "六大 AI 模型适配、React Flow 流程看板、智能去重、Top-N 选优、质量感知、GB/T 7714 引用输出等关键能力均已落地。")
    add_speech(doc, "后续我规划了两个方向, 希望领导指示: "
                    "方向 A 是检索策略深度优化, 例如加入 IPC 分类号联动、引文网络图谱、同族专利智能识别; "
                    "方向 B 是报告系统增强, 例如自动生成审查建议、权利要求对比表自动填充、审查员知识库沉淀。"
                    "这两个方向在技术可行性和业务价值上各有侧重, 具体走哪个, 希望领导给予指示。")

    # --- 第 14 页 结束 ---
    add_page_tag(doc, 14, "结束页 / Q&A", "20 秒")
    add_action_hint(doc, "停顿 2 秒后再开口, 目光扫视全场。")
    add_speech(doc, "最后总结一下: 本项目围绕\u201C提效、提质、可追溯\u201D三大目标, "
                    "构建了一套多模型并行的 AI 辅助专利检索系统, 已实现从上传到报告导出的全流程闭环。"
                    "系统的核心价值, 是让审查员从机械的检索工作中解放出来, 把更多精力投入到创造性的判断与审查中。")
    add_speech(doc, "以上是我的汇报内容, 欢迎各位领导指正与提问。谢谢!")

    # ============== 附录 ==============
    doc.add_page_break()
    add_heading(doc, "附录: 节奏与现场建议", level=1)

    add_heading(doc, "一、整体节奏", level=2)
    add_para(doc, text="\u2022 全文口播约 9 分 30 秒, 留 30 秒缓冲用于切页与视频播放。")
    add_para(doc, text="\u2022 视频段落 (第 5 页) 占用约 2 分 40 秒, 是全程最长的单项。")
    add_para(doc, text="\u2022 方案介绍 (第 6\u20139 页) 每页约 50 秒, 注意控制不要超时。")

    add_heading(doc, "二、可压缩项 (如果时间紧张)", level=2)
    add_para(doc, text="\u2022 第 13 页\u201C后续规划\u201D可压缩到 25 秒, 只说两个方向名称, 等领导追问再展开。")
    add_para(doc, text="\u2022 第 14 页\u201C收尾\u201D可压缩到 15 秒, 直接说核心金句 + 谢谢。")
    add_para(doc, text="\u2022 第 9 页\u201C笛卡尔积\u201D的右侧关键设计可只说前 2 条, 其余一带而过。")

    add_heading(doc, "三、可扩展项 (如果时间富余或领导追问)", level=2)
    add_para(doc, text="\u2022 第 10 页看板可现场打开系统, 演示一次实时任务 (需网络与 API 可用)。")
    add_para(doc, text="\u2022 第 11 页报告可现场切到 DOCX 导出, 展示真实文件。")
    add_para(doc, text="\u2022 第 12 页质量感知可展开讲\u201C为什么不做虚假填充\u201D的设计哲学。")

    add_heading(doc, "四、风险预案", level=2)
    add_para(doc, runs=[
        {"text": "\u2022 视频无法播放: ", "bold": True, "color": ACCENT},
        {"text": "改为口述视频要点 (登录\u2192上传\u2192解析\u2192配置\u2192检索\u2192看板\u2192报告\u2192导出), 请领导会后再看视频。"},
    ])
    add_para(doc, runs=[
        {"text": "\u2022 时间超时: ", "bold": True, "color": ACCENT},
        {"text": "直接跳到第 14 页收尾, 第 13 页只说\u201C已完成核心闭环, 后续两个方向待定\u201D。"},
    ])
    add_para(doc, runs=[
        {"text": "\u2022 领导追问技术细节: ", "bold": True, "color": ACCENT},
        {"text": "可邀请会后单独交流, 或展示 CLAUDE.md / DEVELOPMENT.md 文档。"},
    ])

    # 保存
    doc.save(OUTPUT)
    size_kb = os.path.getsize(OUTPUT) / 1024
    print(f"\u2713 Word 文档已生成: {OUTPUT}")
    print(f"  文件大小: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()